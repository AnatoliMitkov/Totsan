from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs-output/ruse-university-cpp-test.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tbl_cell_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def style_paragraph(paragraph, size=11, color="000000", bold=False):
    for run in paragraph.runs:
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(color)
        run.bold = bold


def add_heading(document, text, level=1):
    paragraph = document.add_paragraph()
    paragraph.style = f"Heading {level}"
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor.from_string("1F4D78")
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(6)
    else:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor.from_string("2E74B5")
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(4)
    return paragraph


def add_body(document, text):
    paragraph = document.add_paragraph(text)
    paragraph.style = "Normal"
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.15
    style_paragraph(paragraph, size=11)
    return paragraph


def add_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.add_run(text)
    paragraph.paragraph_format.space_after = Pt(4)
    style_paragraph(paragraph, size=11)
    return paragraph


def configure_styles(document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color in [
        ("Heading 1", 16, "1F4D78"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)


def configure_page(document):
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def add_footer(document):
    section = document.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Русенски университет | Тестов документ")
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("61758F")


def add_cover(document):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("РУСЕНСКИ УНИВЕРСИТЕТ\n„Ангел Кънчев“")
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    r.font.size = Pt(14)
    r.bold = True

    for _ in range(4):
        document.add_paragraph()

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Кратка разработка по програмиране на C++")
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    r.font.size = Pt(24)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Тема: Основни принципи на обектно-ориентираното програмиране")
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string("1F4D78")

    for _ in range(6):
        document.add_paragraph()

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for line in [
        "Изготвил: студент",
        "Факултет: примерен факултет",
        "Дисциплина: Програмиране на C++",
        "Град: Русе",
        "Година: 2026",
    ]:
        run = meta.add_run(line + "\n")
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        run.font.size = Pt(11)

    document.add_section(WD_SECTION.NEW_PAGE)


def add_comparison_table(document):
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_width(table)
    set_cell_margins(table)
    headers = ["Принцип", "Идея", "Пример в C++"]
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "E8EEF5")
        paragraph = cell.paragraphs[0]
        paragraph.add_run(text)
        style_paragraph(paragraph, size=10, bold=True, color="0B2545")

    rows = [
        ("Капсулация", "Скриване на вътрешното състояние и контролиране на достъпа.", "private полета и public методи."),
        ("Наследяване", "Преизползване и разширяване на поведение от базов клас.", "class Student : public Person."),
        ("Полиморфизъм", "Един интерфейс може да има различни реализации.", "virtual методи и override."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cells[i].paragraphs[0]
            paragraph.add_run(text)
            style_paragraph(paragraph, size=10)


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    configure_page(document)
    configure_styles(document)
    add_footer(document)
    add_cover(document)

    add_heading(document, "1. Въведение", 1)
    add_body(
        document,
        "C++ е език за програмиране с широко приложение в системното програмиране, "
        "инженерния софтуер, игрите и приложенията, при които производителността е важна. "
        "Една от причините езикът да се използва толкова дълго е възможността да съчетава "
        "ниско ниво на контрол с модерни абстракции."
    )
    add_body(
        document,
        "Обектно-ориентираното програмиране позволява програмите да се структурират около "
        "обекти, които обединяват данни и поведение. Това прави по-лесно моделирането на "
        "реални процеси и намалява повторението в кода."
    )

    add_heading(document, "2. Основни принципи", 1)
    add_body(document, "В C++ основните принципи на обектно-ориентираното програмиране могат да се обобщят така:")
    add_bullet(document, "капсулация - контрол върху достъпа до данните в класа;")
    add_bullet(document, "наследяване - създаване на нови класове върху вече съществуващи;")
    add_bullet(document, "полиморфизъм - използване на общ интерфейс за различни реализации.")

    add_heading(document, "3. Сравнителна таблица", 1)
    add_body(document, "Следната таблица показва кратко сравнение между трите принципа и примерна употреба в C++.")
    add_comparison_table(document)

    add_heading(document, "4. Кратък пример", 1)
    add_body(
        document,
        "Ако имаме базов клас Person и производен клас Student, можем да използваме наследяване, "
        "за да избегнем дублирането на общи характеристики като име и възраст. Същевременно "
        "класът Student може да добави специфична информация като факултетен номер или специалност."
    )
    code = document.add_paragraph()
    code.paragraph_format.left_indent = Inches(0.25)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(8)
    run = code.add_run(
        "class Person { ... };\n"
        "class Student : public Person { ... };"
    )
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(10)

    add_heading(document, "5. Заключение", 1)
    add_body(
        document,
        "Обектно-ориентираното програмиране в C++ е важна основа за изграждане на по-големи "
        "и по-поддържани програми. Разбирането на капсулация, наследяване и полиморфизъм "
        "помага на студента да пише по-ясен, по-гъвкав и по-лесен за развитие код."
    )

    add_heading(document, "Използвани източници", 1)
    add_bullet(document, "Bjarne Stroustrup, The C++ Programming Language.")
    add_bullet(document, "cppreference.com - справочник за езика C++.")
    add_bullet(document, "Учебни материали по програмиране на C++.")

    document.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    build()
